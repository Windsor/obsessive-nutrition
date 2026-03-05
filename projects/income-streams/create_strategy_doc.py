#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Multiple Income Streams Strategy', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Prepared by Saturday for Glenn • February 2026')
doc.add_paragraph('')

# Tier 1
doc.add_heading('Tier 1: Fast to Launch, Highly Automatable', level=1)

doc.add_heading('📚 Erotica / Romance Fiction (KDP)', level=2)
doc.add_paragraph('Why: Highest ROI genre on Amazon. Voracious readers, series loyalty, and AI can handle 80% of the writing.')
doc.add_paragraph('Model: Series of 15-20k word novellas, published weekly/biweekly under pen names.')
doc.add_paragraph('Automation: AI writes drafts, you review/edit, publish. Cover generation via AI. Full pipeline management.')
doc.add_paragraph('Revenue: $500-5k/month per active pen name once you have 5-10 titles.')
doc.add_paragraph('Hot niches: dark romance, paranormal romance, reverse harem, taboo, monster romance.')

doc.add_heading('📖 Non-Fiction (KDP)', level=2)
doc.add_paragraph('Why: Evergreen passive income, less competitive than fiction in specific niches.')
doc.add_paragraph('Model: 20-30k word books on specific topics — self-help, business, investing, Portugal/expat life, adult lifestyle guides.')
doc.add_paragraph('Automation: Research + write + publish pipeline, almost fully automatable.')
doc.add_paragraph('Revenue: $200-2k/month per title that finds its niche.')

doc.add_heading('🎬 Faceless YouTube/TikTok Channels', level=2)
doc.add_paragraph('Why: Ad revenue + affiliate links. Faceless = fully automatable.')
p = doc.add_paragraph('Channel ideas:')
doc.add_paragraph('Reddit story narration (text-to-speech + gameplay footage)', style='List Bullet')
doc.add_paragraph('"Dark history" or true crime compilations', style='List Bullet')
doc.add_paragraph('Finance/investing explainers', style='List Bullet')
doc.add_paragraph('Adult lifestyle/relationship advice (monetizes well)', style='List Bullet')
doc.add_paragraph('Portugal travel/expat content', style='List Bullet')
doc.add_paragraph('Automation: Script → AI voiceover → auto-edit → schedule publish.')
doc.add_paragraph('Revenue: $1-10k/month per channel once monetized.')

# Tier 2
doc.add_heading('Tier 2: Medium Setup, Strong Recurring Revenue', level=1)

doc.add_heading('🌐 Directory/Affiliate Sites', level=2)
doc.add_paragraph('Already started with the adult services directory for Portugal.')
doc.add_paragraph('Monetization: Featured listings (€50-200/month), affiliate links, ad revenue.')
doc.add_paragraph('Expand to: Other countries, other niches (cannabis, wellness, nightlife).')
doc.add_paragraph('Automation: Add-listing-via-URL feature already built as foundation.')

doc.add_heading('📧 Newsletter Business', level=2)
doc.add_paragraph('Model: Curated niche newsletters (AI-generated), monetize via sponsorships + paid tier.')
doc.add_paragraph('Niches: Investing, Portugal expat life, adult lifestyle, AI/tech.')
doc.add_paragraph('Automation: AI curates + writes, you approve, auto-send.')
doc.add_paragraph('Revenue: $500-5k/month per newsletter at 5k+ subscribers.')

doc.add_heading('🛠️ SaaS Micro-Tools', level=2)
doc.add_paragraph('Small, focused tools that solve one problem well.')
doc.add_paragraph('Ideas: Adult venue review aggregator, swinger event finder, expat tax calculator for Portugal.')
doc.add_paragraph('Revenue: Subscription-based, $10-50/month per user.')

# Action Plan
doc.add_heading('Action Plan', level=1)

doc.add_heading('Week 1-2', level=2)
doc.add_paragraph('Launch first erotica pen name — 3-book series, publish on KDP', style='List Number')
doc.add_paragraph('Set up one faceless YouTube channel — pick niche, create content pipeline', style='List Number')
doc.add_paragraph('Monetize adult directory — add featured listings, contact venues for paid placement', style='List Number')

doc.add_heading('Week 3-4', level=2)
doc.add_paragraph('Launch non-fiction book #1 (investing, tech, or expat life)', style='List Number')
doc.add_paragraph('Second YouTube channel or start a newsletter', style='List Number')
doc.add_paragraph('Publish books 2-3 of the erotica series', style='List Number')

doc.add_heading('Month 2+', level=2)
doc.add_paragraph('Scale what\'s working, kill what\'s not', style='List Bullet')
doc.add_paragraph('Add more pen names, more channels', style='List Bullet')
doc.add_paragraph('Build SaaS tool if clear opportunity emerges', style='List Bullet')

doc.save('/Users/jarvis/.openclaw/workspace/projects/income-streams/income-strategy.docx')
print('Done')
