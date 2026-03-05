#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

with open('manuscript.md', 'r') as f:
    content = f.read()

lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    # Book title
    if line.startswith('# ') and i < 5:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = 'Georgia'
        i += 1
        continue
    
    # Subtitle/series
    if line.startswith('### ') and i < 10:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[4:])
        run.italic = True
        run.font.size = Pt(14)
        run.font.name = 'Georgia'
        i += 1
        continue
    
    # Chapter headings
    if line.startswith('## '):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(72)
        p.space_after = Pt(24)
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Georgia'
        i += 1
        continue
    
    # POV headers
    if line.startswith('### '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(18)
        run = p.add_run(line[4:])
        run.font.size = Pt(12)
        run.font.name = 'Georgia'
        run.italic = True
        i += 1
        continue
    
    # Horizontal rule / scene break
    if line == '---':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(12)
        p.space_after = Pt(12)
        run = p.add_run('⁕  ⁕  ⁕')
        run.font.size = Pt(11)
        i += 1
        continue
    
    # Regular paragraph - handle markdown italics/bold
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.3)
    
    # Parse inline formatting
    text = line
    # Split by bold and italic markers
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = p.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
    
    i += 1

doc.save('Cruel_Obsession_-_The_Obsession_Duet_Book_1.docx')
print('Done')
